import gradio as gr
from openai import OpenAI
import os
from fpdf import FPDF
import docx

ACCESS_TOKEN = os.getenv("HF_TOKEN")

client = OpenAI(
    base_url="https://api-inference.huggingface.co/v1/",
    api_key=ACCESS_TOKEN,
)

css = '''
.gradio-container{max-width: 1000px !important}
h1{text-align:center}
footer {
    visibility: hidden
}
'''

def respond(
    message,
    history: list[tuple[str, str]],
    system_message,
    max_tokens,
    temperature,
    top_p,
):
    messages = [{"role": "system", "content": system_message}]

    for val in history:
        if val[0]:
            messages.append({"role": "user", "content": val[0]})
        if val[1]:
            messages.append({"role": "assistant", "content": val[1]})

    messages.append({"role": "user", "content": message})

    response = ""
    
    for message in client.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-70B-Instruct",
        max_tokens=max_tokens,
        stream=True,
        temperature=temperature,
        top_p=top_p,
        messages=messages,
    ):
        token = message.choices[0].delta.content
        response += token
        yield response

def save_file(content, filename, file_format, font_name, font_size, alignment):
    font_path = f"font/{font_name}"
    
    if file_format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_font("CustomFont", "", font_path, uni=True)
        pdf.set_font("CustomFont", size=int(font_size))
        
        pdf.set_font_size(int(font_size))
        pdf.ln(1.0 * int(font_size))  # Default line spacing

        align = 'L'
        if alignment == "Center":
            align = 'C'
        elif alignment == "Right":
            align = 'R'
        elif alignment == "Justify":
            align = '' 

        for line in content.split("\n"):
            if alignment == "Justify":
                pdf.multi_cell(0, 10, line)
            else:
                pdf.multi_cell(0, 10, line, align=align)
        
        pdf.output(f"{filename}.pdf")
        return f"{filename}.pdf"
    elif file_format == "docx":
        doc = docx.Document()
        paragraph = doc.add_paragraph(content)
        run = paragraph.add_run(content)
        run.font.name = font_name.split(".")[0]
        run.font.size = docx.shared.Pt(int(font_size)) 
        
        if alignment == "Left":
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "Center":
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "Right":
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        else:  # Justify
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.save(f"{filename}.docx")
        return f"{filename}.docx"
    elif file_format == "txt":
        with open(f"{filename}.txt", "w") as f:
            f.write(content)
        return f"{filename}.txt"
    else:
        raise ValueError("Unsupported file format")

def respond_and_save(message, history, system_message, filename="output", file_format="pdf", font_name="arial.ttf", font_size="18", alignment="Left", max_tokens=512, temperature=0.7, top_p=0.95):
    generated_text = ""
    for output in respond(message, history, system_message, max_tokens, temperature, top_p):
        generated_text = output
    saved_file = save_file(generated_text, filename, file_format, font_name, font_size, alignment)
    return generated_text, history + [(message, generated_text)], saved_file

font_choice = gr.Dropdown(
    choices=[
        "DejaVuMathTeXGyre.ttf", 
        "FiraCode-Medium.ttf", 
        "InputMono-Light.ttf",
        "JetBrainsMono-Thin.ttf", 
        "ProggyCrossed Regular Mac.ttf", 
        "SourceCodePro-Black.ttf", 
        "arial.ttf", 
        "calibri.ttf", 
        "mukta-malar-extralight.ttf", 
        "noto-sans-arabic-medium.ttf", 
        "times new roman.ttf",
        "ANGSA.ttf",
        "Book-Antiqua.ttf",
        "CONSOLA.TTF",
        "COOPBL.TTF",
        "Rockwell-Bold.ttf",
        "Candara Light.TTF",
        "Carlito-Regular.ttf",
        "Castellar.ttf",
        "Courier New.ttf",
        "LSANS.TTF",
        "Lucida Bright Regular.ttf",
        "TRTempusSansITC.ttf",
        "Verdana.ttf",
        "bell-mt.ttf",
        "eras-itc-light.ttf",
        "fonnts.com-aptos-light.ttf",
        "georgia.ttf",
        "segoeuithis.ttf",
        "youyuan.TTF",
        "TfPonetoneExpanded-7BJZA.ttf",
    ],
    value="mukta-malar-extralight.ttf",
    label="Font Style"
)

font_size = gr.Dropdown(
    choices=["6", "8", "10","12", "14", "16", "18", "20", "22", "24"],
    value="12",
    label="Font Size"
)

alignment_choice = gr.Dropdown(
    choices=["Left", "Center", "Right", "Justify"],
    value="Justify",
    label="Text Alignment"
)

demo = gr.Interface(
    fn=respond_and_save,
    inputs=[
        gr.Textbox(placeholder="Type your message here...", label="Chatbot", lines=3),
        gr.State(value=[]),  
        gr.Textbox(placeholder="System message", label="System message", value="", visible=False),
        gr.Textbox(placeholder="Filename (default: output)", label="Filename", value="output", visible=False),
        gr.Radio(["pdf", "docx", "txt"], label="File Format", value="pdf"),
        font_choice,
        font_size,
        alignment_choice,
        gr.Slider(minimum=1, maximum=2048, value=512, step=1, label="Max new tokens", visible=False),
        gr.Slider(minimum=0.1, maximum=4.0, value=0.7, step=0.1, label="Temperature", visible=False),
        gr.Slider(minimum=0.1, maximum=1.0, value=0.95, step=0.05, label="Top-P", visible=False),
    ],
    outputs=[
        gr.Textbox(label="Generated Text", lines=11),
        gr.State(value=[]), 
        gr.File(label="Download File")
    ],
    css=css,
    title="GRABDOC-V 🌵",
    theme="bethecloud/storj_theme"
)
demo.queue().launch(show_api=False)